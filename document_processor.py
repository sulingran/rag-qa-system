import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from config import CHUNK_SIZE, CHUNK_OVERLAP


# ────────────────────────────── parsers ──────────────────────────────

def _parse_pdf(file_path: str, filename: str) -> list[Document]:
    """Extract text per page from a PDF."""
    try:
        import pdfplumber
        docs = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(
                        page_content=text.strip(),
                        metadata={"source": filename, "page": page_num},
                    ))
        return docs
    except Exception:
        # Fallback to pypdf
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        docs = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    page_content=text.strip(),
                    metadata={"source": filename, "page": page_num},
                ))
        return docs


def _parse_txt(file_path: str, filename: str) -> list[Document]:
    """Read entire TXT file as one document."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [Document(page_content=text.strip(), metadata={"source": filename})]


def _parse_docx(file_path: str, filename: str) -> list[Document]:
    """Extract paragraphs from a DOCX file."""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return [Document(page_content=text.strip(), metadata={"source": filename})]


# ────────────────────────────── splitter ─────────────────────────────

def _split(docs: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
    )
    chunks = splitter.split_documents(docs)
    # Tag each chunk with its sequential index within this batch
    for i, chunk in enumerate(chunks, start=1):
        chunk.metadata["chunk_index"] = i
    return chunks


# ────────────────────────────── public API ───────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".doc"}


def process_document(file_path: str, filename: str) -> list[Document]:
    """Parse a document file and return chunked LangChain Documents."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        raw_docs = _parse_pdf(file_path, filename)
    elif ext == ".txt":
        raw_docs = _parse_txt(file_path, filename)
    elif ext in (".docx", ".doc"):
        raw_docs = _parse_docx(file_path, filename)
    else:
        raise ValueError(f"不支持的文件类型: {ext}，请上传 PDF / TXT / DOCX")

    if not raw_docs:
        raise ValueError("文档内容为空，无法处理")

    return _split(raw_docs)
